from docx.api import Document
from docx import Document
import pandas as pd
import matplotlib.pyplot as plt

imgname = 'data.png'
students = pd.read_excel('./data/people.xlsx')

students.sort_values(by='Score', inplace=True, ascending=False)
# students.reset_index(drop=True, inplace=True)
# print('原始students:', students.Name[0])
print('原始students:', students.iloc[0, :]['Name'])

# generate chart

plt.bar(students.Name, students.Score, color='orange')

plt.title('Student Score', fontsize=16)
plt.xlabel('Name')
plt.ylabel('Score')

plt.xticks(students.Name, rotation='90')
plt.tight_layout()
# plt.show()
plt.savefig(imgname)

# generate word
document = Document()
document.add_heading('Data analysis Report', level=0)
first_student = students.iloc[0, :]['Name']
first_score = students.iloc[0, :]['Score']

p = document.add_paragraph('The highest score is student ')
p.add_run(str(first_student)).bold = True
p.add_run(', score is ')
p.add_run(str(first_score)).bold = True

p1 = document.add_paragraph(
    f'Totally {len(students.Name)} students attended the test, the summary of the testing is:')

table = document.add_table(rows=len(students.Name)+1, cols=2)
table.style = 'LightShading-Accent1'

table.cell(0, 0).text = 'Student Name'
table.cell(0, 1).text = 'Student Score'

for i, (index, row) in enumerate(students.iterrows()):
    table.cell(i+1, 0).text = str(row['Name'])
    table.cell(i+1, 1).text = str(row['Score'])

document.add_picture(imgname)
document.save('Students.docx')
print('Done!!')
